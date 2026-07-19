"""Extract plain text from .docx files (inspections, SOPs, incidents, compliance docs)."""
from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P


def extract_docx(filepath: str) -> str:
    """Return all paragraph and table text in TRUE document order.

    python-docx's doc.paragraphs and doc.tables are separate flat lists that
    lose interleaving -- a table appearing near the top of the document (e.g.
    a metadata table with the doc's own ID) would otherwise get extracted
    after body paragraphs that appear later in the actual document (e.g. a
    "References" section mentioning other doc IDs). That ordering matters
    downstream, since entity extraction uses first-occurrence position to
    decide which doc ID is "this document's own ID". Walking doc.element.body
    directly preserves the real order.
    """
    doc = Document(filepath)
    parts = []

    for child in doc.element.body.iterchildren():
        if isinstance(child, CT_P):
            text = Paragraph(child, doc).text.strip()
            if text:
                parts.append(text)
        elif isinstance(child, CT_Tbl):
            table = Table(child, doc)
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                row_text = " | ".join(c for c in cells if c)
                if row_text:
                    parts.append(row_text)

    return "\n".join(parts)
